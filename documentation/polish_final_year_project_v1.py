from __future__ import annotations

from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


DOC_PATH = Path("/Users/eniholla/Desktop/Projects/hotel-management-system/documentation/FINAL YEAR PROJECT V1.docx")
BACKUP_PATH = DOC_PATH.with_name("FINAL YEAR PROJECT V1.pre_polish_backup.docx")


def para(doc: Document, number: int):
    return doc.paragraphs[number - 1]


def set_text(paragraph, text: str) -> None:
    paragraph.text = text


def set_style(paragraph, style_name: str) -> None:
    paragraph.style = style_name


def insert_paragraph_after(paragraph, text: str = "", style_name: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    if style_name:
        new_paragraph.style = paragraph.part.document.styles[style_name]
    return new_paragraph


def add_toc_field(paragraph: Paragraph) -> None:
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Update the table to display the final contents."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)


def enable_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)


def ensure_required_styles(doc: Document) -> None:
    if "Heading 1" not in [style.name for style in doc.styles]:
        heading_one = doc.styles.add_style("Heading 1", WD_STYLE_TYPE.PARAGRAPH, builtin=True)
        heading_one.base_style = doc.styles["Heading 2"]
        heading_one.priority = 9
        heading_one.quick_style = True
        heading_one.unhide_when_used = True
    if "Title" not in [style.name for style in doc.styles]:
        title = doc.styles.add_style("Title", WD_STYLE_TYPE.PARAGRAPH)
        title.base_style = doc.styles["Body Text"]


def main() -> None:
    if not DOC_PATH.exists():
        raise FileNotFoundError(DOC_PATH)

    if not BACKUP_PATH.exists():
        shutil.copy2(DOC_PATH, BACKUP_PATH)

    doc = Document(DOC_PATH)
    ensure_required_styles(doc)

    original_paragraphs = list(doc.paragraphs)

    text_updates = {
        6: "Design and Implementation of an IoT-Enabled Online Hotel Management and Reservation System",
        27: "We declare that this project work, titled Design and Implementation of an IoT-Enabled Online Hotel Management and Reservation System, was carried out by the following students:",
        55: "We hereby certify that this project was carried out by the underlisted students under the supervision of the Department of Computer Science, School of Computing, Babcock University, Ilishan-Remo, Ogun State, Nigeria.",
        76: "This project is dedicated first to God Almighty, whose grace, strength, and guidance sustained us throughout the course of this work. We also dedicate it to our supervisor, Dr. Olawale Somefun, and to everyone whose encouragement and support helped us bring the project to completion.",
        98: "We sincerely appreciate everyone who contributed to the successful completion of this project. Above all, we thank God Almighty for His grace, mercy, and guidance throughout the course of this work. We are deeply grateful to our project supervisor, Dr. Olawale Somefun, for his patient guidance, constructive feedback, and encouragement. We also appreciate our parents, family members, friends, classmates, and lecturers for their support, prayers, and motivation. Every form of help we received played an important part in bringing this project to completion.",
        118: "The hospitality industry depends on timely reservation handling, accurate room-status control, secure payment processing, and coordinated staff workflow. This project presents the design and implementation of an IoT-enabled online hotel management and reservation system developed to improve hotel operations while strengthening health, safety, and environmental awareness. The system was built with Django, HTML, CSS, and JavaScript, and it supports room browsing, online and offline booking, Paystack-backed payment processing, role-based access control, custom administrative and receptionist dashboards, housekeeping coordination, and simulated IoT room monitoring. The monitoring module tracks conditions such as temperature, gas level, and motion activity, stores sensor readings, and generates alerts when abnormal situations are detected. The completed system demonstrates how hotel administration and smart monitoring can be integrated into a single platform that improves operational efficiency, accountability, and guest safety. Overall, the project provides a practical prototype for smarter hotel management and a strong foundation for future real-world deployment.",
        119: "Keywords: hotel management system, IoT monitoring, reservation system, Paystack payment, role-based access control, HSE alerting.",
        149: "TABLE OF CONTENTS",
        164: "CHAPTER ONE: INTRODUCTION",
        177: "The aim of this project is to design and implement an IoT-enabled online hotel management and reservation system that improves hotel operations, strengthens record management, supports secure booking and payment workflows, and enhances safety through room-condition monitoring.",
        181: "1.1 Problem Statement",
        197: "1.2 Significance of the Study",
        214: "1.3 Aim and Objectives",
        233: "This study focuses on the design and implementation of an online hotel management and reservation system with integrated IoT-style monitoring. The system covers room browsing, online and offline booking, role-based dashboards, payment processing, customer record management, housekeeping coordination, and room-condition alerts based on simulated sensor data. The work does not extend to full smart-building automation, multi-branch deployment, or physical hardware installation at production scale. Its emphasis is on a functional academic prototype that demonstrates both hotel operations and safety-aware monitoring in one platform.",
        236: "The system was implemented with HTML, CSS, JavaScript, and Django-based templates on the frontend, while Python and Django were used for backend logic, routing, model management, and authentication. SQLite served as the prototype database, Paystack was integrated for online payment verification, and the IoT component was implemented through simulated monitoring services that generate room-condition data for dashboard display and alert handling. This combination made it possible to demonstrate the intended workflow end to end within a realistic academic prototype.",
        237: "The development process followed the major SDLC stages of requirement analysis, system design, implementation, testing, and documentation.",
        253: "CHAPTER TWO: LITERATURE REVIEW",
        254: "",
        255: "2.0 Overview of Hotel Management Systems",
        285: "CHAPTER THREE: SYSTEM ANALYSIS AND DESIGN",
        286: "",
        292: "In developing this software, an iterative approach guided by Agile principles was adopted. This model was suitable because the project combines several connected modules, including guest booking, room management, payment processing, receptionist operations, housekeeping support, and IoT-based monitoring. Working in iterations made it possible to build the core workflow first and then refine the dashboards, alert logic, and supporting operational features as the system matured.",
        293: "",
        294: "This approach aligns well with the changing demands of hotel software. By dividing development into manageable stages, the team could focus on one feature area at a time, gather feedback, test progress, and improve the system continuously without losing sight of the overall objective.",
        301: "3.2 Requirement Analysis",
        308: "The architecture follows a standard layered Django pattern. The presentation layer is formed by HTML templates, Bootstrap-driven layout elements, and page-specific CSS or JavaScript enhancements. The application layer contains the view logic, validation rules, role checks, session handling, and feature orchestration. The data layer is provided by Django models backed by SQLite for the prototype environment.",
        313: "3.4 Unified Modelling Language (UML) Diagrams",
        317: "3.4.1 Use Case Diagram",
        332: "3.4.2 Activity Diagram",
        346: "3.5 Data Design",
        357: "3.6 Workflow Design",
        361: "3.7 Development Tools and Environment",
        375: "4. Integrated Development Environment (IDE): Visual Studio Code was used as the main development environment for coding, debugging, and project organization. It supported a smoother workflow across the Django application, templates, and supporting scripts used in the implementation.",
        392: "CHAPTER FOUR: SYSTEM IMPLEMENTATION AND RESULTS",
        393: "",
        399: "The system was developed using a layered architecture made up of the following components:",
        400: "• Frontend Layer - HTML, CSS, JavaScript, and Django templates for user interaction.",
        401: "• Application Layer - Django views, forms, models, and services handling business logic.",
        402: "• Database Layer - SQLite for prototype storage of users, bookings, payments, tasks, and sensor data.",
        403: "• Monitoring Layer - Python-based simulation and alert services for room-condition updates.",
        404: "• Integration Layer - Paystack callback handling and configurable email/SMS notification support.",
        407: "• User authentication and role-based access control.",
        408: "• Room inventory, online booking, and offline booking management.",
        409: "• Paystack payment verification workflow.",
        410: "• Administrative dashboard and reporting views.",
        411: "• Receptionist operations and housekeeping coordination.",
        412: "• IoT monitoring and HSE alert management.",
        475: "",
        579: "4.6.1 Motion (PIR) Logic",
        588: "4.6.2 Smoke/Gas Detection (MQ-2) Logic",
        600: "4.6.3 Temperature Sensor Logic",
        613: "4.6.4 Simulation Code Overview",
        688: "CHAPTER FIVE: SUMMARY, CONCLUSION AND RECOMMENDATIONS",
        689: "",
        694: "This project successfully delivered an IoT-enabled web-based hotel management and reservation system that combines guest booking, administrative control, receptionist workflow, payment handling, and room-condition monitoring in one platform.",
        696: "The system brings together core hotel operations and smart monitoring by providing:",
        705: "The project achieved its main objective by producing a practical, reliable, and academically relevant platform for hotel management and safety-aware operations.",
        713: "It demonstrates that software engineering and IoT-oriented monitoring can work together to improve efficiency, accountability, and situational awareness in hospitality environments.",
        724: "Future improvements may include:",
        735: "This project contributes to knowledge by:",
        745: "5.6 References",
        769: "Management System. ResearchGate. Available at 10.54097/fcis.v4i2.9866",
    }

    for number, text in text_updates.items():
        set_text(original_paragraphs[number - 1], text)

    heading_one_indices = [25, 54, 74, 97, 117, 164, 253, 285, 392, 688]
    heading_two_indices = [
        167, 181, 197, 214, 231, 235, 240,
        255, 261, 264, 267, 270, 276, 279,
        288, 301, 307, 313, 346, 357, 361,
        395, 414, 522, 556, 564, 574, 616, 631, 638, 661, 671,
        692, 703, 716, 722, 733, 745,
    ]
    heading_three_indices = [
        200, 203, 206, 209,
        290, 317, 332,
        398, 406, 416, 433, 446, 458, 484, 498, 508,
        530, 538, 540, 550, 579, 588, 600, 613, 640, 644, 648, 652, 656,
    ]

    for index in heading_one_indices:
        set_style(original_paragraphs[index - 1], "Heading 1")
        original_paragraphs[index - 1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for index in heading_two_indices:
        set_style(original_paragraphs[index - 1], "Heading 2")
        original_paragraphs[index - 1].alignment = WD_ALIGN_PARAGRAPH.LEFT

    for index in heading_three_indices:
        set_style(original_paragraphs[index - 1], "Heading 3")
        original_paragraphs[index - 1].alignment = WD_ALIGN_PARAGRAPH.LEFT

    toc_heading = original_paragraphs[149 - 1]
    set_style(toc_heading, "Title")
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_heading.paragraph_format.space_after = 0

    abstract_heading = original_paragraphs[117 - 1]
    abstract_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    toc_paragraph = insert_paragraph_after(toc_heading)
    add_toc_field(toc_paragraph)

    overview_intro = insert_paragraph_after(
        original_paragraphs[255 - 1],
        "A hotel is a commercial establishment that provides paid lodging on a short-term basis. Hotels play an important role in the tourism and hospitality industry by creating jobs, supporting travel, generating revenue, and contributing to local economic activity.",
        style_name="Body A",
    )
    overview_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    enable_update_fields_on_open(doc)
    doc.save(DOC_PATH)

    print(f"Updated {DOC_PATH}")
    print(f"Backup available at {BACKUP_PATH}")


if __name__ == "__main__":
    main()
