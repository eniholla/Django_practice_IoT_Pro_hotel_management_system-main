from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from tempfile import TemporaryDirectory
import textwrap

from PIL import Image
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT_DIR = Path("/Users/eniholla/Desktop/Projects/hotel-management-system")
DOC_DIR = ROOT_DIR / "documentation"
OUTPUT_DOCX = DOC_DIR / "chapter_2_3_4_upscaled_project_documentation.docx"
OUTPUT_PREVIEW = DOC_DIR / "chapter_2_3_4_upscaled_preview.md"
ASSET_DIR = DOC_DIR / "assets"
SCREENSHOT_DIR = ASSET_DIR / "screenshots"
FLOWCHART_DIR = ASSET_DIR / "flowcharts"
DIAGRAM_DIR = ASSET_DIR / "diagrams"
ERD_DIAGRAM = DIAGRAM_DIR / "hotel_system_erd.png"


@dataclass(frozen=True)
class Snippet:
    title: str
    source: Path
    start_line: int
    end_line: int
    note: str

    @property
    def source_label(self) -> str:
        return f"{self.source.relative_to(ROOT_DIR)} (lines {self.start_line}-{self.end_line})"


@dataclass(frozen=True)
class FeatureSection:
    title: str
    overview: list[str]
    implementation_note: str
    screenshot: Path
    screenshot_caption: str
    flowchart: Path
    flowchart_caption: str
    snippets: list[Snippet]


CHAPTER2_SECTIONS = [
    (
        "2.0 Introduction",
        [
            "Literature review is important in a final-year computing project because it shows where a proposed system belongs within the wider body of practice and prior design thinking. For this study, the review is not limited to reservation software in isolation. The project sits at the intersection of hotel operations, role-based information systems, digital payment processing, and smart-room monitoring.",
            "The present hotel management system was developed as a practical answer to familiar operational problems: fragmented room records, slow front-desk coordination, weak visibility into payment state, and poor linkage between ordinary hotel transactions and health, safety, and environment awareness. A useful review therefore has to examine both conventional hotel management software and the newer expectation that software should support operational intelligence rather than record keeping alone.",
        ],
    ),
    (
        "2.1 Evolution of Hotel Management Systems",
        [
            "Hotel operations were once managed through paper ledgers, wall charts, manual cashbooks, and face-to-face front-desk communication. Those methods were workable for small guest houses, but they became unreliable as hotels needed faster room turnover, better audit trails, and clearer visibility into who was arriving, departing, or owing payment. Manual systems also made it difficult to coordinate between reception, management, and support staff in real time.",
            "The first wave of computerized hotel tools focused mainly on digitizing bookings, room allocation, and billing. Later web-based systems expanded that scope by allowing public room browsing, remote reservation, and centralized staff access. In the current stage of development, hospitality software is expected to do more than store data. It is increasingly judged by how well it supports live decisions, role-specific dashboards, secure payment confirmation, and operational monitoring across the property.",
        ],
    ),
    (
        "2.2 Core Functions in Contemporary Hotel Applications",
        [
            "A modern hotel platform usually combines several responsibilities inside one workflow. It must manage room inventory, customer identity, booking windows, check-in and check-out events, payment state, and staff activity. These modules are tightly connected. A booking affects room status. A payment event affects confirmation state. A checkout event affects housekeeping readiness. Good software reflects these dependencies instead of treating each activity as a disconnected form.",
            "Another major expectation in current systems is role separation. Guests need simple room discovery and reservation pages. Administrators need a strategic view of rooms, users, employees, revenue-related figures, and alerts. Reception staff need a narrow but faster operational view that emphasizes today's arrivals, departures, room readiness, guest search, and pending payments. When those views are not separated properly, the system becomes cluttered and slows down the users who depend on it most.",
        ],
    ),
    (
        "2.3 Smart Hotel Monitoring and HSE Awareness",
        [
            "Smart-hotel thinking adds a second layer to ordinary hotel management: visibility into what is happening physically inside the building. Temperature drift, gas leakage, unusual motion, or device silence may all be operationally important even when the reservation records look normal. This matters because a room can be correctly marked as occupied or available while still presenting an environmental or security risk that staff should respond to quickly.",
            "In many academic and prototype environments, real sensors are not always available. Even so, simulated monitoring remains valuable when it is designed responsibly. A simulation can still model the logic of device registration, payload generation, threshold evaluation, alert persistence, acknowledgement, and resolution. In that sense, a simulation-driven IoT layer is still useful for demonstrating how HSE monitoring can be embedded into hotel workflow rather than left as a theoretical add-on.",
        ],
    ),
    (
        "2.4 Observed Gaps in Existing Approaches",
        [
            "A recurring weakness in many hotel software examples is that the guest-facing booking experience is improved while the internal operating experience remains shallow. Systems often allow room reservation, yet provide limited support for receptionist-specific actions such as fast guest lookup, guided check-in and check-out, housekeeping follow-up, or monitoring of unsettled payments. That gap matters in practice because reception is where operational delays become visible first.",
            "A second gap is the weak connection between payment confirmation and booking persistence. In many simplified prototypes, payment is discussed conceptually but not handled as an actual control point in the reservation lifecycle. A third gap is the separation of safety monitoring from normal hotel software. When room-condition alerts live outside the same application that manages bookings and staff actions, accountability becomes weaker and response time may suffer.",
        ],
    ),
    (
        "2.5 Relevance to the Present Project",
        [
            "The current project responds directly to those gaps by combining public room discovery, structured online booking, Paystack-backed payment verification, administrative control, receptionist workflow support, and an IoT-style monitoring module inside one Django application. This makes the system stronger as a project document because the chapters can discuss a single integrated platform rather than unrelated feature demonstrations.",
            "The project is also relevant as a smart-hospitality prototype. It shows that even without physical hardware devices, a system can still preserve the end-to-end monitoring pattern: assign a logical device to each room, generate or receive readings, evaluate them against room context, raise alerts, and expose those alerts on operational dashboards. That integration is the main conceptual strength of the work and is what distinguishes it from a basic booking website.",
        ],
    ),
    (
        "2.6 Conclusion",
        [
            "The literature and practice background reviewed in this chapter point to a clear direction: hotel software is most useful when it connects reservations, payments, staff workflow, and operational visibility in one coherent system. This project follows that direction. The next chapter therefore explains how the system was designed and implemented as a structured web application with both transactional and monitoring responsibilities.",
        ],
    ),
]

CHAPTER2_COMPARISON_ROWS = [
    ["Dimension", "Traditional hotel handling", "Ordinary web hotel tool", "Present project"],
    ["Reservation capture", "Desk ledger or phone call", "Online form only", "Online form plus validated booking workflow"],
    ["Payment linkage", "Manual reconciliation", "Sometimes external to reservation logic", "Paystack verification tied to booking confirmation"],
    ["Staff access", "Verbal coordination", "Often one generic admin view", "Separated admin and receptionist dashboards"],
    ["Room turnover visibility", "Paper notes or memory", "Partial status visibility", "Room state and housekeeping status stored together"],
    ["Safety awareness", "Manual observation", "Usually absent", "Simulated IoT monitoring with persistent alerts"],
]

CHAPTER3_SECTIONS = [
    (
        "3.0 Introduction",
        [
            "This chapter explains how the hotel management system was analysed, structured, and prepared for implementation. The focus is not only on software coding, but on how the project was broken into coherent modules that reflect real hotel work: user access, room inventory, booking validation, payment progression, receptionist operations, and room-condition monitoring.",
            "Because the system serves different categories of users and combines both transactional logic and simulated environmental monitoring, the design had to be modular. The methodology described here therefore emphasizes layered architecture, explicit data models, and feature flows that can be tested independently while still contributing to one unified application.",
        ],
    ),
    (
        "3.1 Development Approach",
        [
            "An iterative development approach was adopted. The project did not emerge as one large implementation pass; it matured feature by feature. Public pages and authentication came first, followed by room inventory and booking logic, then payment control, staff dashboards, housekeeping linkage, and finally the IoT alerting workflow. This sequence reduced complexity because each stage built on a stable foundation from the previous one.",
            "The iterative approach was especially useful because it allowed the project to remain realistic. For example, the payment flow was upgraded from a simple record-keeping idea into a Paystack-backed confirmation step, and the receptionist workflow evolved into a dedicated dashboard rather than remaining a generic admin extension. The same pattern applied to the HSE component, which matured from a concept into a database-backed monitoring subsystem.",
        ],
    ),
    (
        "3.2 Requirement Analysis",
        [
            "Requirement gathering was guided by the actual roles represented in the system. Guests require visibility into available rooms, stay cost, booking dates, and account access. Administrators require deeper control over rooms, bookings, users, employees, salaries, and monitoring summaries. Reception staff require quick-action tools for arrivals, departures, search, payments, and housekeeping coordination. The monitoring layer, though automated, behaves as another operational stakeholder because it continuously contributes room-condition events to the system.",
            "Functional requirements were therefore grouped around role-specific outcomes rather than abstract menus. The system had to redirect users to the correct dashboard, prevent room clashes, preserve payment state, support room-status updates, create housekeeping tasks after checkout, and evaluate simulated room-condition data against meaningful thresholds. Non-functional requirements were equally important: usability, role isolation, persistence, and maintainability were all necessary for the project to remain credible as a final-year system.",
        ],
    ),
    (
        "3.3 System Architecture",
        [
            "The architecture follows a standard layered Django pattern. The presentation layer is formed by HTML templates, Bootstrap-driven layout elements, and page-specific CSS or JavaScript enhancements. The application layer contains the view logic, validation rules, role checks, session handling, and feature orchestration. The data layer is provided by Django models backed by SQLite for the prototype environment.",
            "The architecture is deliberately split across two application domains. The `HotelApp` module handles users, rooms, bookings, payments, employees, salaries, housekeeping tasks, and activity logs. The `alerts` module manages logical devices, sensor readings, alert lifecycle, notification records, and monitoring snapshots. This split keeps the smart-monitoring responsibilities close to the hotel workflow without allowing them to overwhelm the ordinary reservation code.",
        ],
    ),
    (
        "3.4 Data Design",
        [
            "The data design of the project mirrors the operational relationships inside a hotel. Room records sit at the centre because both online and offline bookings depend on them, payments refer back to them indirectly through booking type and booking identifier, and housekeeping tasks are created around their turnover state. The custom user model adds role information that makes it possible to route administrators, reception staff, and guests differently after authentication.",
            "For the monitoring subsystem, additional entities were introduced instead of overloading the ordinary hotel tables. Logical device records, sensor readings, room-condition alerts, and notification logs are persisted separately so that the monitoring lifecycle can be audited over time. That separation keeps the design clean and allows the HSE module to remain expandable if physical devices are later connected to the platform.",
        ],
    ),
    (
        "3.5 Workflow Design",
        [
            "The public guest workflow begins with room browsing, continues into the booking form, validates stay dates and guest counts, stores a pending reservation in session, and then moves to the payment page. Only after payment verification does the system create the persistent booking record and mark the room as reserved. That sequence is important because it avoids the common prototype problem where bookings are written before payment outcome is known.",
            "The internal operational workflow is different. Administrators use the dashboard as a control surface for aggregated information and management links, while reception staff use a more focused workflow around arrivals, departures, room status, housekeeping, and desk payments. The IoT workflow runs alongside both human paths. It generates or receives room-condition data, evaluates that data, writes sensor history, and updates alerts that become visible to staff through the monitoring pages.",
        ],
    ),
    (
        "3.6 Development Environment and Tools",
        [
            "The implemented stack is centered on Python and Django. HTML, CSS, JavaScript, and template rendering support the interface. SQLite provides a lightweight but reliable persistence layer for the academic prototype. Paystack is used as the active payment gateway in the guest booking flow, while optional notification backends allow the monitoring module to record or dispatch alert messages.",
            "For the documentation refresh itself, a local demo database was generated so that screenshots, dashboards, and alert states would be internally consistent. This was necessary because a presentation-ready report should not rely on accidental live data. Instead, the figures shown in Chapter 4 are derived from a controlled demo environment built directly from the repository's current models and views.",
        ],
    ),
    (
        "3.7 Summary",
        [
            "This chapter has described the design logic behind the project: iterative development, role-aware requirements, modular architecture, and workflow-led implementation. With that design foundation established, Chapter 4 presents the completed features as they exist in the running system and discusses how they behave in a realistic demonstration context.",
        ],
    ),
]

CHAPTER3_STAKEHOLDER_ROWS = [
    ["Stakeholder", "Primary responsibility in the system", "Main interface focus"],
    ["Guest", "Browse rooms, submit reservations, review booking/payment state", "Public pages and booking flow"],
    ["Administrator", "Monitor hotel-wide performance and manage operational records", "Custom admin dashboard"],
    ["Receptionist", "Handle arrivals, departures, search, desk payment, and room readiness", "Reception dashboard and boards"],
    ["Monitoring service", "Generate and evaluate room-condition readings", "IoT dashboard and alert center"],
]

CHAPTER3_ENTITY_ROWS = [
    ["Entity", "Purpose in the implemented system"],
    ["Authorregis", "Custom email-based user model with receptionist and staff roles"],
    ["Room", "Stores room number, type, floor, facilities, price, availability, and housekeeping state"],
    ["OnlineBooking", "Persists guest-submitted reservations after successful payment"],
    ["OfflineBooking", "Persists front-desk or walk-in bookings"],
    ["Employee", "Stores hotel staff bio-data and department records used in operations"],
    ["Salary", "Stores compensation records linked to each employee entry"],
    ["Payment", "Tracks payment method, status, receipt, and Paystack references through booking_type and booking_id"],
    ["HousekeepingTask", "Coordinates room-turnover work after checkout or support actions"],
    ["ActivityLog", "Records important operational actions for audit and visibility"],
    ["IoTDevice", "Assigns a logical monitoring device to each room"],
    ["SensorReading", "Stores temperature, gas, motion, and overall room condition history"],
    ["RoomConditionAlert", "Stores active and resolved HSE incidents tied to readings"],
    ["AlertNotification", "Stores simulated email or SMS notices generated from room-condition alerts"],
]

FEATURES = [
    FeatureSection(
        title="4.2.1 Login and Role-Based Access",
        overview=[
            "The login module gives the system its operational shape because it determines what kind of interface the user sees after authentication. A guest should not land on a crowded control panel, and a receptionist should not be forced to navigate the same screen as a system administrator. The application therefore uses role-aware routing immediately after successful sign-in.",
            "From a presentation standpoint, this feature matters because it is the entry point to every other workflow shown in the document. It demonstrates that the project is not a single flat website, but a structured application with clear access boundaries and user-sensitive navigation.",
        ],
        implementation_note="The implementation keeps the routing logic compact by using helper functions that choose a safe `next` destination or fall back to the correct dashboard for the current role.",
        screenshot=SCREENSHOT_DIR / "login_page.png",
        screenshot_caption="Login interface showing the email-based access form used to enter the hotel system.",
        flowchart=FLOWCHART_DIR / "login_flow.png",
        flowchart_caption="Role-aware login flow from credential submission to dashboard redirection.",
        snippets=[
            Snippet(
                title="Role-aware redirect helper",
                source=ROOT_DIR / "Django_practice_Pro_hotel_management_system-main/HotelApp/views.py",
                start_line=47,
                end_line=70,
                note="This helper centralizes post-login routing and keeps redirection decisions predictable for guests, staff, and reception users.",
            ),
        ],
    ),
    FeatureSection(
        title="4.2.2 Rooms Catalogue and Room Management Context",
        overview=[
            "The rooms feature is more than a gallery. It is the system's public inventory layer, and every booking, payment, dashboard summary, and room-status decision depends on it. The room records store pricing, room type, operational status, facilities, and housekeeping information, which means the same entity serves both guest-facing and staff-facing needs.",
            "In the user experience, the room catalogue translates that stored information into a clear visual list. In the internal workflow, the same room model supports reservation checks, dashboard statistics, maintenance visibility, and turnover readiness.",
        ],
        implementation_note="The room model was designed with both commercial and operational fields so that one record can support public display, booking validation, and internal control views without duplicated data structures.",
        screenshot=SCREENSHOT_DIR / "rooms_page.png",
        screenshot_caption="Public room catalogue showing room type, price, facilities, and current availability state.",
        flowchart=FLOWCHART_DIR / "rooms_flow.png",
        flowchart_caption="Room-discovery flow from public catalogue browsing to booking handoff.",
        snippets=[
            Snippet(
                title="Core room model",
                source=ROOT_DIR / "Django_practice_Pro_hotel_management_system-main/HotelApp/models.py",
                start_line=57,
                end_line=89,
                note="This model defines the inventory attributes that drive both the catalogue view and the staff dashboards.",
            ),
        ],
    ),
    FeatureSection(
        title="4.2.3 Booking Workflow",
        overview=[
            "The booking workflow was implemented as a controlled sequence rather than a form that immediately creates a reservation record. The guest selects a room, chooses dates, enters occupancy details, and triggers server-side validation that checks date order and room-conflict conditions. Only when the booking data is valid does the system move it forward to the payment stage.",
            "That approach improves reliability. It ensures that stay details are preserved long enough to support payment processing, yet the final booking record is not written prematurely. For a final-year project, this is a more realistic workflow than simply saving any submitted reservation request without transactional control.",
        ],
        implementation_note="The system stores the validated booking payload in session as `pending_booking`, which creates a clean handoff between booking validation and Paystack payment.",
        screenshot=SCREENSHOT_DIR / "booking_page.png",
        screenshot_caption="Booking page showing room context, date selection, and guest-count inputs before payment.",
        flowchart=FLOWCHART_DIR / "booking_flow.png",
        flowchart_caption="Validated booking flow from stay details to the pending-payment handoff.",
        snippets=[
            Snippet(
                title="Booking validation and session handoff",
                source=ROOT_DIR / "Django_practice_Pro_hotel_management_system-main/HotelApp/views.py",
                start_line=521,
                end_line=558,
                note="This excerpt shows how the view validates user input and stores a pending booking in session before payment begins.",
            ),
        ],
    ),
    FeatureSection(
        title="4.2.4 Paystack Payment Workflow",
        overview=[
            "Payment is an active part of the booking lifecycle in the current implementation. The guest reviews a structured payment summary and then proceeds through Paystack checkout. The application does not simply assume success; it initializes a transaction, stores a pending payment record, waits for the callback, verifies the returned reference, and only then creates the confirmed booking.",
            "This feature significantly strengthens the realism of the project. It demonstrates how a hospitality system can connect booking data, payment state, activity logging, and room reservation status inside one workflow. It also shows that the platform has evolved beyond the earlier academic-prototype pattern of internal payment notes only.",
        ],
        implementation_note="The payment flow is split into two controlled stages: transaction initialization and callback verification. This allows booking creation to depend on server-side confirmation rather than on client-side assumption.",
        screenshot=SCREENSHOT_DIR / "payment_page.png",
        screenshot_caption="Payment summary page used to review reservation details before redirecting to Paystack.",
        flowchart=FLOWCHART_DIR / "payment_flow.png",
        flowchart_caption="Paystack-backed payment sequence from pending booking to verified reservation.",
        snippets=[
            Snippet(
                title="Paystack transaction initialization",
                source=ROOT_DIR / "Django_practice_Pro_hotel_management_system-main/HotelApp/views.py",
                start_line=1414,
                end_line=1466,
                note="The initialization stage calculates the payable amount, creates a unique reference, and opens a pending payment record before redirecting to Paystack.",
            ),
            Snippet(
                title="Payment callback and booking confirmation",
                source=ROOT_DIR / "Django_practice_Pro_hotel_management_system-main/HotelApp/views.py",
                start_line=1486,
                end_line=1523,
                note="The callback stage verifies the reference, creates the final booking, updates the payment status, and reserves the room.",
            ),
        ],
    ),
    FeatureSection(
        title="4.2.5 Administrative Dashboard",
        overview=[
            "The custom administrative dashboard acts as the management cockpit of the system. It brings together booking totals, room availability, occupancy rate, user counts, salary figures, booking trends, and recent monitoring alerts inside a single page. This gives administrators a working picture of the hotel's current state instead of forcing them to inspect isolated tables one by one.",
            "The dashboard is especially important in this project because it also bridges ordinary hotel operations with smart monitoring. The HSE summary appears in the same space as the business metrics, which reinforces the idea that safety is an operational concern rather than a side system.",
        ],
        implementation_note="The dashboard view aggregates counts, trend values, revenue estimates, and IoT summaries directly from the database so that the page behaves like a live operational dashboard rather than a static mock-up.",
        screenshot=SCREENSHOT_DIR / "admin_dashboard.png",
        screenshot_caption="Custom administrator dashboard showing operational summaries and integrated HSE monitoring cards.",
        flowchart=FLOWCHART_DIR / "admin_flow.png",
        flowchart_caption="Administrative control flow from dashboard load to live management actions.",
        snippets=[
            Snippet(
                title="Administrative metric aggregation",
                source=ROOT_DIR / "Django_practice_Pro_hotel_management_system-main/HotelApp/views.py",
                start_line=177,
                end_line=233,
                note="This excerpt shows how the dashboard compiles room, booking, user, revenue, and IoT summary values before rendering the page.",
            ),
        ],
    ),
    FeatureSection(
        title="4.2.6 Receptionist Dashboard and Front-Desk Operations",
        overview=[
            "The receptionist module was designed around day-to-day front-desk work rather than broad system administration. The dashboard emphasizes what matters during live operations: today's check-ins, today's check-outs, room statistics, pending housekeeping, recent activity, revenue collected today, and pending payments. This keeps the screen practical for front-desk decision making.",
            "The workflow continues into related pages such as room status, guest search, housekeeping, check-in, check-out, and desk payment processing. In other words, the receptionist dashboard is not merely a smaller admin dashboard. It is a role-specific workspace tuned for speed, visibility, and frequent actions.",
        ],
        implementation_note="The receptionist view computes operational metrics around the current date and combines reservation queues with room status and payment summaries so that staff can act from one compact page.",
        screenshot=SCREENSHOT_DIR / "receptionist_dashboard.png",
        screenshot_caption="Receptionist dashboard focused on arrivals, departures, room status, and front-desk quick actions.",
        flowchart=FLOWCHART_DIR / "receptionist_flow.png",
        flowchart_caption="Front-desk workflow from sign-in through room, guest, and housekeeping actions.",
        snippets=[
            Snippet(
                title="Receptionist dashboard aggregation",
                source=ROOT_DIR / "Django_practice_Pro_hotel_management_system-main/HotelApp/views.py",
                start_line=963,
                end_line=1038,
                note="This view builds the arrival/departure queues, room statistics, and payment summaries that make the receptionist interface task-oriented.",
            ),
        ],
    ),
    FeatureSection(
        title="4.2.7 IoT Monitoring and HSE Alerting",
        overview=[
            "The IoT monitoring module is the most distinctive part of the project. Each room is mapped to a logical device, sensor readings are stored over time, and rules convert those readings into warning or critical conditions based on temperature, gas level, motion state, and expected occupancy. The result is a monitoring dashboard that behaves like a live operational subsystem rather than a decorative front-end panel.",
            "The HSE value of the feature lies in persistence and lifecycle. The system does not merely flash a warning; it stores readings, opens alerts, allows acknowledgement, records resolution, and makes recent incidents visible to administrators. That history is important because hotel safety work depends on accountability, not on momentary visual effects.",
        ],
        implementation_note="The monitoring logic separates payload evaluation from persistence. That design keeps the thresholds explicit and allows the same alert lifecycle to work whether readings are simulated or later connected to real devices.",
        screenshot=SCREENSHOT_DIR / "iot_dashboard.png",
        screenshot_caption="IoT monitoring dashboard showing room-condition summaries and active HSE alert counts.",
        flowchart=FLOWCHART_DIR / "iot_flow.png",
        flowchart_caption="Monitoring lifecycle from sensor payload evaluation to alert refresh on the admin pages.",
        snippets=[
            Snippet(
                title="Room-condition rule evaluation",
                source=ROOT_DIR / "Django_practice_Pro_hotel_management_system-main/alerts/services.py",
                start_line=213,
                end_line=311,
                note="The rule engine evaluates temperature, gas, and motion signals against room context to classify each reading as normal, warning, or critical.",
            ),
            Snippet(
                title="Sensor persistence and alert synchronization",
                source=ROOT_DIR / "Django_practice_Pro_hotel_management_system-main/alerts/services.py",
                start_line=314,
                end_line=335,
                note="Once evaluated, each payload is stored as a sensor reading and immediately synchronized with the active alert set for the room.",
            ),
        ],
    ),
]

CHAPTER4_TEST_SECTIONS = [
    "System validation was approached from two directions. First, the project was exercised manually through the same demonstration paths shown in this document: login, room browsing, booking entry, payment summary display, admin dashboard loading, receptionist dashboard loading, and IoT dashboard review. Second, the repository's Django test suite was executed to confirm the health of the built-in regression coverage around login routing, dashboard behaviour, receptionist bootstrap, and monitoring logic.",
    "The automated test run executed 35 tests. Most of the suites completed successfully, particularly the routing, home-view resilience, receptionist environment bootstrap, and IoT monitoring tests. Three older custom-admin assertions failed because they no longer align with the current seeded-room behaviour and current admin form outcomes. That result does not invalidate the documented feature set, but it does show that the admin regression suite should be refreshed as the project evolves.",
]

CHAPTER4_TEST_ROWS = [
    ["Scenario", "Validation style", "Observed result"],
    ["Login redirects by user role", "Automated and manual", "Working in the current codebase and visible in the demo pages"],
    ["Room catalogue displays inventory and status", "Manual", "Working in the demo environment used for documentation screenshots"],
    ["Booking form preserves validated pending data", "Manual and code inspection", "Working through session handoff to the payment summary page"],
    ["Payment flow initializes Paystack transaction data", "Code inspection and manual summary-page review", "Implemented in the active project flow"],
    ["Admin dashboard compiles live metrics", "Manual and automated coverage", "Visible in the demo environment; some older admin tests need maintenance"],
    ["Reception dashboard builds daily queues and summaries", "Automated and manual", "Working in both the rendered demo page and repository logic"],
    ["IoT monitoring records readings and alerts", "Automated and manual", "Working in the repository test suite and the generated demo dashboard"],
]

CHAPTER4_DISCUSSION = [
    "Taken together, the implemented features show that the project has moved beyond a narrow booking website. It now behaves like a compact hospitality operations platform with payment control, staff workflow support, and HSE visibility. The most important design achievement is integration: reservation state, room state, payment state, staff actions, and alert state all live inside one application.",
    "The documentation refresh also makes that integration easier to communicate during presentation. Each feature is shown with its interface, logic flow, and code evidence, which means the final-year report can speak to both technical depth and practical usability. That balance is often where academic system projects struggle, and it is the main reason this regenerated document is more suitable for presentation than the earlier draft.",
]


def require_assets() -> None:
    required = [feature.screenshot for feature in FEATURES] + [feature.flowchart for feature in FEATURES] + [ERD_DIAGRAM]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing documentation assets:\n" + "\n".join(missing))


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for style_name, size in [("Heading 1", 18), ("Heading 2", 15), ("Heading 3", 13), ("Heading 4", 12)]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    if "CaptionStyle" not in document.styles:
        caption = document.styles.add_style("CaptionStyle", WD_STYLE_TYPE.PARAGRAPH)
        caption.font.name = "Times New Roman"
        caption.font.size = Pt(10.5)
        caption.font.italic = True
        caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def set_paragraph_format(paragraph, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(8)) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = space_after
    paragraph.paragraph_format.line_spacing = 1.15


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    set_paragraph_format(paragraph)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None, size: float = 11.0) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_table(document: Document, rows: list[list[str]], *, caption: str | None = None) -> None:
    if caption:
        caption_paragraph = document.add_paragraph(style="CaptionStyle")
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.add_run(caption)

    table = document.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for col_index, value in enumerate(row):
            set_cell_text(cells[col_index], value, bold=row_index == 0)
            if row_index == 0:
                shade_cell(cells[col_index], "D9EAF7")
    document.add_paragraph()


def optimize_image(path: Path, temp_dir: Path) -> Path:
    image = Image.open(path)
    max_size = (1800, 1800) if path.parent == DIAGRAM_DIR else (1300, 1300)
    image.thumbnail(max_size)
    if path.parent in {FLOWCHART_DIR, DIAGRAM_DIR}:
        output = temp_dir / f"{path.stem}.png"
        image.save(output, format="PNG", optimize=True)
    else:
        output = temp_dir / f"{path.stem}.jpg"
        image = image.convert("RGB")
        image.save(output, format="JPEG", quality=82, optimize=True)
    return output


def add_figure(
    document: Document,
    image_path: Path,
    caption: str,
    temp_dir: Path,
    *,
    width_inches: float,
    figure_label: str,
) -> None:
    optimized = optimize_image(image_path, temp_dir)
    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(optimized), width=Inches(width_inches))

    caption_paragraph = document.add_paragraph(style="CaptionStyle")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.add_run(f"{figure_label}: {caption}")


def extract_snippet(snippet: Snippet) -> str:
    lines = snippet.source.read_text(encoding="utf-8").splitlines()
    block = "\n".join(lines[snippet.start_line - 1:snippet.end_line])
    return textwrap.dedent(block).strip("\n")


def add_code_block(document: Document, snippet: Snippet) -> None:
    note = document.add_paragraph()
    note_run = note.add_run(f"Code extract: {snippet.title}. ")
    note_run.bold = True
    note_run.font.name = "Times New Roman"
    note_run.font.size = Pt(11)
    note_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    tail = note.add_run(f"{snippet.note} Source: {snippet.source_label}.")
    tail.italic = True
    tail.font.name = "Times New Roman"
    tail.font.size = Pt(10.5)
    tail._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    set_paragraph_format(note, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(4))

    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    shade_cell(table.cell(0, 0), "F4F7FB")
    cell = table.cell(0, 0)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(extract_snippet(snippet).splitlines()):
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9.2)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        if index != len(extract_snippet(snippet).splitlines()) - 1:
            run.add_break()
    document.add_paragraph()


def add_chapter_heading(document: Document, title: str) -> None:
    heading = document.add_paragraph(style="Heading 1")
    heading_run = heading.add_run(title)
    heading_run.font.name = "Times New Roman"
    heading_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_feature_section(document: Document, feature: FeatureSection, temp_dir: Path, figure_index: int) -> int:
    document.add_paragraph(feature.title, style="Heading 3")
    for paragraph in feature.overview:
        add_body_paragraph(document, paragraph)

    implementation_intro = document.add_paragraph()
    label = implementation_intro.add_run("Implementation note: ")
    label.bold = True
    label.font.name = "Times New Roman"
    label.font.size = Pt(12)
    label._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    body = implementation_intro.add_run(feature.implementation_note)
    body.font.name = "Times New Roman"
    body.font.size = Pt(12)
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    set_paragraph_format(implementation_intro)

    add_figure(
        document,
        feature.screenshot,
        feature.screenshot_caption,
        temp_dir,
        width_inches=5.6,
        figure_label=f"Figure 4.{figure_index}",
    )
    figure_index += 1
    add_figure(
        document,
        feature.flowchart,
        feature.flowchart_caption,
        temp_dir,
        width_inches=6.0,
        figure_label=f"Figure 4.{figure_index}",
    )
    figure_index += 1

    for snippet in feature.snippets:
        add_code_block(document, snippet)

    return figure_index


def build_docx() -> None:
    require_assets()
    document = Document()
    configure_styles(document)

    with TemporaryDirectory() as temp_name:
        temp_dir = Path(temp_name)

        add_chapter_heading(document, "Chapter 2")
        for title, paragraphs in CHAPTER2_SECTIONS:
            document.add_paragraph(title, style="Heading 2")
            for paragraph in paragraphs:
                add_body_paragraph(document, paragraph)
            if title == "2.4 Observed Gaps in Existing Approaches":
                add_table(
                    document,
                    CHAPTER2_COMPARISON_ROWS,
                    caption="Table 2.1: Comparison of traditional hotel handling, ordinary web tools, and the present integrated project.",
                )

        document.add_page_break()

        add_chapter_heading(document, "Chapter 3")
        for title, paragraphs in CHAPTER3_SECTIONS:
            document.add_paragraph(title, style="Heading 2")
            for paragraph in paragraphs:
                add_body_paragraph(document, paragraph)
            if title == "3.2 Requirement Analysis":
                add_table(
                    document,
                    CHAPTER3_STAKEHOLDER_ROWS,
                    caption="Table 3.1: Stakeholder groups and the interfaces that serve them in the implemented system.",
                )
            if title == "3.4 Data Design":
                add_table(
                    document,
                    CHAPTER3_ENTITY_ROWS,
                    caption="Table 3.2: Core persistent entities and their responsibilities in the current implementation.",
                )
                add_figure(
                    document,
                    ERD_DIAGRAM,
                    "Entity-relationship diagram showing how the booking, operations, payment, and IoT monitoring tables are connected in the current implementation.",
                    temp_dir,
                    width_inches=6.3,
                    figure_label="Figure 3.1",
                )
                add_body_paragraph(
                    document,
                    "Figure 3.1 complements Table 3.2 by showing the actual structural links behind the system. Solid connectors represent direct foreign-key or one-to-one relationships in the Django models, while the dashed payment connector highlights the current application's logical booking link implemented through `booking_type` and `booking_id` rather than a strict database foreign key.",
                )

        document.add_page_break()

        add_chapter_heading(document, "Chapter 4")
        document.add_paragraph("4.0 Introduction", style="Heading 2")
        add_body_paragraph(
            document,
            "This chapter presents the implemented system in a presentation-ready format. Rather than describing the project only in abstract terms, the chapter ties each major feature to the live interface, the underlying control flow, and selected source-code excerpts from the current repository. The result is a chapter that can support both viva presentation and technical defence.",
        )
        add_body_paragraph(
            document,
            "The figures in this chapter were generated from a controlled demo dataset built from the repository's current models and views. That approach keeps the screenshots internally consistent and allows the documented dashboards, booking states, and alert conditions to reflect a coherent operating scenario instead of random leftover data.",
        )

        document.add_paragraph("4.1 Implemented Feature Overview", style="Heading 2")
        add_body_paragraph(
            document,
            "The completed application is best understood as a connected hospitality platform made of seven visible feature areas: access control, room inventory, booking, payment, administration, reception workflow, and HSE monitoring. Each one is presented below with the same evidence pattern so that the reader can move naturally from interface to logic.",
        )

        document.add_paragraph("4.2 Feature Walkthrough", style="Heading 2")
        add_body_paragraph(
            document,
            "The subsections below present each major feature with the same structure: functional explanation, implementation note, interface evidence, process flow, and a short code extract from the current repository.",
        )

        figure_index = 1
        for feature in FEATURES:
            figure_index = add_feature_section(document, feature, temp_dir, figure_index)

        document.add_paragraph("4.3 Testing and Validation", style="Heading 2")
        for paragraph in CHAPTER4_TEST_SECTIONS:
            add_body_paragraph(document, paragraph)
        add_table(
            document,
            CHAPTER4_TEST_ROWS,
            caption="Table 4.1: Summary of the main validation scenarios considered during the documentation refresh.",
        )

        document.add_paragraph("4.4 Discussion", style="Heading 2")
        for paragraph in CHAPTER4_DISCUSSION:
            add_body_paragraph(document, paragraph)

        document.add_paragraph("4.5 Summary", style="Heading 2")
        add_body_paragraph(
            document,
            "This chapter has shown the final system as it exists today: role-aware, payment-backed, operationally structured, and enriched with simulated IoT monitoring. The presentation-focused layout of this refreshed documentation now makes the project easier to defend academically because each major feature is supported by interface evidence, flow logic, and implementation excerpts from the real codebase.",
        )

        OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
        document.save(OUTPUT_DOCX)


def add_md_paragraph(lines: list[str], text: str) -> None:
    lines.append(text)
    lines.append("")


def add_md_table(lines: list[str], rows: list[list[str]], *, caption: str | None = None) -> None:
    if caption:
        lines.append(f"*{caption}*")
        lines.append("")
    header = rows[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * len(header)) + " |")
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")


def add_md_feature(lines: list[str], feature: FeatureSection, figure_index: int) -> int:
    lines.append(f"### {feature.title}")
    lines.append("")
    for paragraph in feature.overview:
        add_md_paragraph(lines, paragraph)

    add_md_paragraph(lines, f"**Implementation note:** {feature.implementation_note}")
    lines.append(f"![{feature.screenshot_caption}]({feature.screenshot.relative_to(DOC_DIR).as_posix()})")
    lines.append("")
    lines.append(f"*Figure 4.{figure_index}: {feature.screenshot_caption}*")
    lines.append("")
    figure_index += 1
    lines.append(f"![{feature.flowchart_caption}]({feature.flowchart.relative_to(DOC_DIR).as_posix()})")
    lines.append("")
    lines.append(f"*Figure 4.{figure_index}: {feature.flowchart_caption}*")
    lines.append("")
    figure_index += 1

    for snippet in feature.snippets:
        lines.append(f"**Code extract: {snippet.title}**")
        lines.append("")
        lines.append(f"*{snippet.note} Source: `{snippet.source_label}`.*")
        lines.append("")
        lines.append("```python")
        lines.append(extract_snippet(snippet))
        lines.append("```")
        lines.append("")

    return figure_index


def build_markdown() -> None:
    lines: list[str] = []
    lines.append("# Chapter 2")
    lines.append("")
    for title, paragraphs in CHAPTER2_SECTIONS:
        lines.append(f"## {title}")
        lines.append("")
        for paragraph in paragraphs:
            add_md_paragraph(lines, paragraph)
        if title == "2.4 Observed Gaps in Existing Approaches":
            add_md_table(
                lines,
                CHAPTER2_COMPARISON_ROWS,
                caption="Table 2.1: Comparison of traditional hotel handling, ordinary web tools, and the present integrated project.",
            )

    lines.append("# Chapter 3")
    lines.append("")
    for title, paragraphs in CHAPTER3_SECTIONS:
        lines.append(f"## {title}")
        lines.append("")
        for paragraph in paragraphs:
            add_md_paragraph(lines, paragraph)
        if title == "3.2 Requirement Analysis":
            add_md_table(
                lines,
                CHAPTER3_STAKEHOLDER_ROWS,
                caption="Table 3.1: Stakeholder groups and the interfaces that serve them in the implemented system.",
            )
        if title == "3.4 Data Design":
            add_md_table(
                lines,
                CHAPTER3_ENTITY_ROWS,
                caption="Table 3.2: Core persistent entities and their responsibilities in the current implementation.",
            )
            lines.append("![Entity-relationship diagram for the hotel management system](assets/diagrams/hotel_system_erd.png)")
            lines.append("")
            lines.append("*Figure 3.1: Entity-relationship diagram showing how the booking, operations, payment, and IoT monitoring tables are connected in the current implementation.*")
            lines.append("")
            add_md_paragraph(
                lines,
                "Figure 3.1 complements Table 3.2 by showing the actual structural links behind the system. Solid connectors represent direct foreign-key or one-to-one relationships in the Django models, while the dashed payment connector highlights the current application's logical booking link implemented through `booking_type` and `booking_id` rather than a strict database foreign key.",
            )

    lines.append("# Chapter 4")
    lines.append("")
    lines.append("## 4.0 Introduction")
    lines.append("")
    add_md_paragraph(
        lines,
        "This chapter presents the implemented system in a presentation-ready format. Rather than describing the project only in abstract terms, the chapter ties each major feature to the live interface, the underlying control flow, and selected source-code excerpts from the current repository. The result is a chapter that can support both viva presentation and technical defence.",
    )
    add_md_paragraph(
        lines,
        "The figures in this chapter were generated from a controlled demo dataset built from the repository's current models and views. That approach keeps the screenshots internally consistent and allows the documented dashboards, booking states, and alert conditions to reflect a coherent operating scenario instead of random leftover data.",
    )

    lines.append("## 4.1 Implemented Feature Overview")
    lines.append("")
    add_md_paragraph(
        lines,
        "The completed application is best understood as a connected hospitality platform made of seven visible feature areas: access control, room inventory, booking, payment, administration, reception workflow, and HSE monitoring. Each one is presented below with the same evidence pattern so that the reader can move naturally from interface to logic.",
    )

    lines.append("## 4.2 Feature Walkthrough")
    lines.append("")
    add_md_paragraph(
        lines,
        "The subsections below present each major feature with the same structure: functional explanation, implementation note, interface evidence, process flow, and a short code extract from the current repository.",
    )

    figure_index = 1
    for feature in FEATURES:
        figure_index = add_md_feature(lines, feature, figure_index)

    lines.append("## 4.3 Testing and Validation")
    lines.append("")
    for paragraph in CHAPTER4_TEST_SECTIONS:
        add_md_paragraph(lines, paragraph)
    add_md_table(
        lines,
        CHAPTER4_TEST_ROWS,
        caption="Table 4.1: Summary of the main validation scenarios considered during the documentation refresh.",
    )

    lines.append("## 4.4 Discussion")
    lines.append("")
    for paragraph in CHAPTER4_DISCUSSION:
        add_md_paragraph(lines, paragraph)

    lines.append("## 4.5 Summary")
    lines.append("")
    add_md_paragraph(
        lines,
        "This chapter has shown the final system as it exists today: role-aware, payment-backed, operationally structured, and enriched with simulated IoT monitoring. The presentation-focused layout of this refreshed documentation now makes the project easier to defend academically because each major feature is supported by interface evidence, flow logic, and implementation excerpts from the real codebase.",
    )

    OUTPUT_PREVIEW.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def main() -> None:
    build_docx()
    build_markdown()
    print(OUTPUT_DOCX)
    print(OUTPUT_PREVIEW)


if __name__ == "__main__":
    main()
